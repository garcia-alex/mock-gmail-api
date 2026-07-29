"""Render parsed section content into a Word document using python-docx."""

from __future__ import annotations

import io

from docx import Document

from mock_gmail_api.docgen.textblocks import parse_sections


def render(title: str, text: str) -> bytes:
    document = Document()
    document.add_heading(title, level=0)

    for section in parse_sections(text):
        if section.heading:
            document.add_heading(section.heading, level=1)
        for block in section.blocks:
            if block.kind == "bullets" and block.items:
                for item in block.items:
                    document.add_paragraph(item, style="List Bullet")
            elif block.text:
                document.add_paragraph(block.text)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
